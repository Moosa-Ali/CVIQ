"""CVIQ smoke test — canonical-JSON / content-presentation architecture.

Single flat ``test_smoke()``, first-failure abort, prints
``ALL SMOKE TESTS PASSED``. Green against the FINAL architecture:

- Canonical JSON only (CVData + CustomSection), Suggestion.field_path /
  rationale, ConfidenceFlag, JDProfile, GapDiff.
- Uploads: parser.parse_file (pdf/docx, heuristic), /api/cv/parse adds
  classification + confidence_flags, sessions store both.
- Rendering is deterministic (no AI): render_html (Jinja2) -> fitz.Story
  html_to_pdf -> reportlab-legacy fallback only; docx_export renders
  custom_sections. Gallery templates are metadata-only (render_template /
  converted / previews); NO slot detection, NO /api/export/preserved, no
  preserve/overlay/template_fill modules.
- Gap engine: deterministic_checks (rules, no AI) + gap_analysis
  {jd_profile, deterministic, semantic, mode}; jd_parser LLM+fallback;
  classify heuristic/merge flags.

Run from the project root:
    $env:PYTHONPATH = "C:\\D-Drive\\Work\\CVIQ"; & ".\\.venv\\Scripts\\python.exe" tests\\smoke_test.py
"""

import base64
import io
import json
import os
import re
import tempfile
import zipfile

import docx
from docx.shared import Pt
import fitz  # noqa: F401  (PyMuPDF; fitz deprecation warning on import is acceptable)
from fastapi.testclient import TestClient

# ---------------------------------------------------------------------------
# Isolate app state BEFORE importing app.*: a developer's committed
# .cvmod/config.json (or env overrides) would make /api/cv/parse and
# /api/cv/gaps call a real LLM. Point CVMOD_DATA_DIR at a fresh temp dir and
# clear the LLM env overrides so the suite is 100% deterministic and offline.
# ---------------------------------------------------------------------------
_TEST_DATA_DIR = tempfile.mkdtemp(prefix="cviq-smoke-")
os.environ["CVMOD_DATA_DIR"] = _TEST_DATA_DIR
for _env in ("OPENROUTER_API_KEY", "AWS_ACCESS_KEY_ID", "AWS_SECRET_ACCESS_KEY", "AWS_REGION"):
    os.environ.pop(_env, None)

from app.main import app
from app.services import templates as template_service
from app.services.cv import analyzer as cv_analyzer
from app.services.cv import classify as cv_classify
from app.services.cv import gaps as cv_gaps
from app.services.cv import jd_parser as cv_jd_parser
from app.services.cv import markdown as cv_markdown
from app.services.cv import session as cv_session
from app.services.cv import tailor as cv_tailor
from app.services.cv.models import CVData, CustomSection, ExperienceItem, Suggestion
from app.services.cv.parser import parse_file
from app.services.export.docx_export import export_docx
from app.services.export.pdf_export import export_pdf
from app.services.export.pdf_render import html_to_pdf, render_pdf
from app.services.export.render_html import render_html
from app.routes import cv as cv_routes


def roundPct(v):
    """Round a 0..1 heuristic confidence to an integer percentage (mirrors the
    frontend helper so parse-confidence smoke checks read the same way)."""
    return int(round(float(v) * 100.0))


# ---------------------------------------------------------------------------
# FakeLLM — dispatches on prompt substrings. Branch ORDER matters: several
# prompts share substrings ("exact shape" appears in classify, JD parse and
# semantic-gaps prompts; '"suggestions"' appears inside the analyze schema), so
# the most specific markers win first.
# ---------------------------------------------------------------------------
CANNED_ANALYSIS = json.dumps(
    {
        "ats_score": 82,
        "matched_keywords": [{"keyword": "python", "count": 4}],
        "missing_keywords": [{"keyword": "kubernetes", "count": 0}],
        "sections": [{"section": "summary", "score": 90, "comment": "good"}],
        "comments": ["looks solid"],
        "gaps": ["Add Kubernetes experience", "No containerization keywords"],
        "suggestions": [
            {
                "id": "v2s-1",
                "section": "experience",
                "field": "bullets",
                "index": 0,
                "type": "rewrite",
                "title": "Quantify impact",
                "original": "Built widgets",
                "suggested": "Built widgets for 200 users",
                "reason": "add metrics",
                "priority": "high",
                "impact": "ATS keyword match",
                "move_from": "experience",
                "move_to": "summary",
                "target_section": "skills",
                "field_path": "experience[0].bullets[0]",
                "rationale": "Metrics raise ATS scoring.",
            },
            {
                "id": "v2s-2",
                "section": "experience",
                "field": "bullets",
                "index": None,
                "type": "rename",
                "title": "Use standard ATS heading",
                "original": "Employment",
                "suggested": "Experience",
                "reason": "Standard ATS heading is safer",
                "priority": "medium",
                "impact": "ATS heading match",
                "field_path": "summary",
                "rationale": "Standard headings are more ATS-friendly.",
            },
        ],
    }
)

# Generic-mode analysis: matched/missing MUST be empty (no JD to match against).
CANNED_GENERIC_ANALYSIS = json.dumps(
    {
        "ats_score": 78,
        "matched_keywords": [],
        "missing_keywords": [],
        "sections": [{"section": "summary", "score": 70, "comment": "ok"}],
        "comments": ["generic ATS-friendliness review"],
        "gaps": ["No quantified achievements", "Consider reordering sections"],
        "suggestions": [],
    }
)

CANNED_SUGGESTIONS = json.dumps(
    {
        "suggestions": [
            {
                "id": "v2g-1",
                "section": "experience",
                "field": "bullets",
                "index": 0,
                "type": "reword",
                "title": "Quantify impact",
                "original": "helped build stuff",
                "suggested": "led delivery of",
                "reason": "Add metrics",
                "priority": "high",
                "impact": "ATS keyword match",
                "field_path": "experience[0].bullets[0]",
                "rationale": "Quantified bullets score higher with ATS and recruiters.",
            },
            {
                "id": "v2g-2",
                "section": "summary",
                "field": "summary",
                "index": None,
                "type": "rename",
                "title": "Use standard ATS heading",
                "original": "Employment",
                "suggested": "Experience",
                "reason": "Standard ATS heading is safer",
                "priority": "medium",
                "impact": "ATS heading match",
                "field_path": "summary",
                "rationale": "Standard headings are more ATS-friendly.",
            },
        ]
    }
)

CANNED_JD = json.dumps(
    {
        "role_title": "Senior Python Engineer",
        "required_skills": ["Python", "Kubernetes", "Docker"],
        "nice_to_have_skills": ["AWS"],
        "must_have_keywords": ["python", "kubernetes", "docker"],
        "seniority_signals": ["senior"],
        "requirements": ["5+ years building Python services"],
    }
)

CANNED_SEMANTIC_GAPS = json.dumps(
    {
        "gaps": [
            {
                "field_path": "experience[0].bullets[0]",
                "issue": "Bullet lacks numbers",
                "suggested_value": "",
                "rationale": "Add metrics to the first bullet",
                "severity": "low",
            }
        ]
    }
)

CANNED_CLASSIFICATION = json.dumps(
    {
        "cv": {
            "personal": {"name": "Ada Lovelace", "title": "Software Engineer", "email": "a@x.com"},
            "summary": "Experienced in python.",
            "experience": [
                {"company": "Acme", "role": "Engineer", "location": "",
                 "dates": {"start": "", "end": ""}, "bullets": ["helped build stuff"]}
            ],
            "skills": [{"category": "lang", "skills": ["Python", "SQL"]}],
        },
        "low_confidence": [{"field_path": "personal.phone", "reason": "not detected"}],
    }
)

# Chat 2.0: strict JSON {reply, proposed_edits:[Suggestion]}.
CANNED_CHAT = json.dumps(
    {
        "reply": "I tightened your summary to lead with impact.",
        "proposed_edits": [
            {
                "id": "summary-stronger-1",
                "section": "summary",
                "field": "summary",
                "index": None,
                "type": "rewrite",
                "title": "Strengthen summary",
                "original": "Experienced in python.",
                "suggested": "Senior engineer with deep Python expertise.",
                "reason": "Lead with impact.",
                "priority": "high",
                "impact": "ATS keyword match",
                "field_path": "summary",
                "rationale": "Stronger opening.",
            }
        ],
    }
)


class FakeLLM:
    provider_name = "fake"

    def __init__(self):
        self.calls = 0

    def chat(self, messages, temperature=0.4, max_tokens=2048):
        self.calls += 1
        prompt = messages[-1]["content"]
        if isinstance(prompt, list):  # vision/classification block content
            prompt = " ".join(
                b.get("text", "") for b in prompt if isinstance(b, dict) and b.get("type") == "text"
            )
        if "converting the candidate's CV below" in prompt:
            return CANNED_CLASSIFICATION
        if "job description below into a structured profile" in prompt:
            return CANNED_JD
        if "gap analysis" in prompt:
            return CANNED_SEMANTIC_GAPS
        if "exact shape" in prompt:
            return CANNED_ANALYSIS
        if '"suggestions"' in prompt:
            return CANNED_SUGGESTIONS
        return "Rewritten segment text."

    def test(self):
        return None


class CapturingLLM:
    """FakeLLM that records every chat() call and returns a canned response.

    Keeps the full ``messages`` list plus the last user-message ``content`` so
    tests can assert the canonical content-block convention (text/image blocks)
    and the markdown prompt content ("## " + OBSERVED LAYOUT) end to end.
    """

    provider_name = "fake"

    def __init__(self, response: str):
        self.response = response
        self.messages = None
        self.user_content = None

    def chat(self, messages, temperature=0.4, max_tokens=2048):
        self.messages = messages
        self.user_content = messages[-1]["content"]
        return self.response

    def test(self):
        return None


class RaisingLLM:
    """Deterministic LLM that always fails — used to prove gap analysis and JD
    parsing degrade gracefully (never raise) when the provider misbehaves."""

    provider_name = "raising"

    def chat(self, messages, temperature=0.4, max_tokens=2048):
        from app.services.llm.base import LLMError

        raise LLMError("deterministic test failure")

    def test(self):
        return None


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def make_cv():
    return CVData(
        personal={"name": "Ada Lovelace", "title": "SWE", "email": "a@x.com"},
        summary="Experienced in python.",
        experience=[{"company": "Acme", "role": "Engineer", "bullets": ["helped build stuff"]}],
        skills=[{"category": "lang", "skills": ["Python", "SQL"]}],
    )


def make_cv_without_email():
    cv = make_cv()
    cv.personal.email = ""
    return cv


def make_cv_with_custom_sections():
    cv = make_cv()
    cv.custom_sections = [CustomSection(title="Volunteering", bullets=["Helped at a shelter"])]
    return cv


def make_resume_pdf():
    """Headered text PDF: name/title, bold EXPERIENCE + SKILLS headings, role,
    dates and bullets. Exercises PDF section-heading detection and the
    layout-aware markdown renderer (OBSERVED LAYOUT subsection)."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 90), "Ada Lovelace")
    page.insert_text((72, 120), "Software Engineer")
    page.insert_text((72, 150), "EXPERIENCE", fontname="hebo", fontsize=14)
    page.insert_text((72, 180), "Senior Engineer at Acme")
    page.insert_text((72, 210), "Jan 2020 - Feb 2021")
    page.insert_text((72, 240), "Led the team")
    page.insert_text((72, 270), "Built widgets")
    page.insert_text((72, 300), "SKILLS", fontname="hebo", fontsize=14)
    page.insert_text((72, 330), "Python, Java")
    data = doc.tobytes()
    doc.close()
    return data


def make_scanned_pdf():
    """2-page image-only PDF (vector rectangles, no text) -> is_image_pdf."""
    doc = fitz.open()
    for _ in range(2):
        page = doc.new_page()
        page.draw_rect(fitz.Rect(50, 50, 320, 220), color=(0.1, 0.1, 0.1), width=2)
        page.draw_rect(fitz.Rect(60, 260, 420, 340), color=(0.1, 0.1, 0.1), width=1)
    data = doc.tobytes()
    doc.close()
    return data


def make_text_pdf():
    """Text-based PDF with the lines consumed by the parse/round-trip tests."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 90), "Ada Lovelace")
    page.insert_text((72, 120), "Software Engineer")
    page.insert_text((72, 150), "Python, Java")
    page.insert_text((72, 180), "Led the team")
    data = doc.tobytes()
    doc.close()
    return data


def make_docx_bytes(name="Ada Lovelace"):
    """DOCX fixture with real Word styles + run sizes.

    - name paragraph, name run size Pt 22
    - ``WORK EXPERIENCE`` Heading 1
    - ``Senior Engineer - Acme`` body paragraph
    - ``Led the team`` List Bullet whose run is Pt(10.5)
    - ``SKILLS`` Heading 1
    - ``Python, Java`` body paragraph
    """
    doc = docx.Document()
    p_name = doc.add_paragraph(name)
    p_name.runs[0].font.size = Pt(22)
    doc.add_heading("WORK EXPERIENCE", level=1)
    doc.add_paragraph("Senior Engineer - Acme")
    bullet = doc.add_paragraph("Led the team", style="List Bullet")
    bullet.runs[0].font.size = Pt(10.5)
    doc.add_heading("SKILLS", level=1)
    doc.add_paragraph("Python, Java")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_long_cv(jobs=40, bullets=8):
    """Multi-page CV (40 jobs x 8 bullets) for the pagination regression."""
    cv = make_cv()
    cv.experience = [
        ExperienceItem(
            company=f"Company {i}",
            role=f"Senior Engineer {i}",
            bullets=[f"delivered outcome {j} for milestone {i}" for j in range(bullets)],
        )
        for i in range(jobs)
    ]
    return cv


def test_smoke():
    client = TestClient(app)
    fake = FakeLLM()
    app.dependency_overrides[cv_routes.get_llm_client] = lambda: fake

    created_session_ids = []

    # ------------------------------------------------------------------
    # S1 -- meta + config (get / save / redaction / test-without-creds)
    # ------------------------------------------------------------------
    r = client.get("/api/meta")
    assert r.status_code == 200, r.text
    assert "templates" in r.json()

    r = client.get("/api/config")
    assert r.status_code == 200, r.text
    assert r.json()["configured"] is False, "tests must start with no LLM configured"

    # config/test with no credentials -> LLMConfigError -> 400, no network.
    r = client.post("/api/config/test", json={"provider": "openrouter"})
    assert r.status_code == 400, r.text

    # direct config_store round trip + redaction.
    data_dir = tempfile.mkdtemp(prefix="cviq-cfg-")
    from app.services.llm import config_store
    from app.services.llm.base import LLMConfig

    config_store.save_config(LLMConfig(openrouter_api_key="secret123"), data_dir)
    loaded = config_store.load_config(data_dir)
    assert loaded.openrouter_api_key == "secret123"
    assert loaded.redacted().openrouter_api_key == "***"

    # ------------------------------------------------------------------
    # S2 -- frontend SPA served by FastAPI
    # ------------------------------------------------------------------
    r = client.get("/")
    assert r.status_code == 200
    assert "CV" in r.text or "app" in r.text.lower()
    r = client.get("/app.js")
    assert r.status_code == 200

    # ------------------------------------------------------------------
    # S3 -- validate
    # ------------------------------------------------------------------
    r = client.post("/api/cv/validate", json={"cv": make_cv().model_dump()})
    assert r.status_code == 200, r.text
    assert "warnings" in r.json()

    # ------------------------------------------------------------------
    # S4 -- parse pipeline (parse_file units)
    # ------------------------------------------------------------------
    parsed = parse_file(make_text_pdf(), "sample.pdf")
    assert parsed["structure"]["kind"] == "pdf"
    assert parsed["is_image_pdf"] is False
    assert parsed["page_images"] == []
    assert isinstance(parsed["page_count"], int)
    assert "Ada" in parsed["text"] and parsed["confidence"] > 0
    assert roundPct(parsed["confidence"]) > 0

    headered = parse_file(make_resume_pdf(), "headered.pdf")
    assert headered["cv"].experience, "PDF heading detection failed to populate experience"
    assert headered["cv"].skills, "PDF heading detection failed to populate skills"

    docx_parsed = parse_file(make_docx_bytes(), "cv.docx")
    assert docx_parsed["structure"]["kind"] == "docx"
    assert docx_parsed["is_image_pdf"] is False
    assert docx_parsed["page_count"] == 0
    assert docx_parsed["page_images"] == []
    assert docx_parsed["confidence"] > 0, "DOCX parse must yield a positive confidence"

    # valid-but-empty DOCX is rejected as "no extractable text".
    empty_docx_buf = io.BytesIO()
    docx.Document().save(empty_docx_buf)
    try:
        parse_file(empty_docx_buf.getvalue(), "empty.docx")
        raise AssertionError("expected ValueError for empty DOCX content")
    except ValueError as exc:
        assert "extractable text" in str(exc), str(exc)

    # garbage DOCX bytes are an invalid zip: fail loudly (never half-baked).
    try:
        parse_file(b"whatever", "garbage.docx")
        raise AssertionError("expected failure for garbage DOCX bytes")
    except zipfile.BadZipFile:
        pass

    # everything else (e.g. .txt) is rejected with the dual-format message.
    try:
        parse_file(b"whatever", "notes.txt")
        raise AssertionError("expected ValueError for non-PDF/non-DOCX extension (.txt)")
    except ValueError as exc:
        assert "PDF or DOCX" in str(exc), str(exc)

    # scanned PDF -> image mode + page renders available for the vision path.
    scan = parse_file(make_scanned_pdf(), "scan.pdf")
    assert scan["is_image_pdf"] is True
    assert scan["page_count"] == 2
    assert len(scan["page_images"]) == 2
    assert scan["structure"]["kind"] == "pdf"
    assert scan["confidence"] == 0.0, "scanned CVs must carry zero parse confidence"

    # ------------------------------------------------------------------
    # S5 -- confidence flags (heuristic, no AI)
    # ------------------------------------------------------------------
    flags = cv_classify.heuristic_flags(make_cv_without_email())
    assert flags and any(f.field_path == "personal.email" for f in flags), flags

    # ------------------------------------------------------------------
    # S6 -- markdown rendering (structure_to_markdown + document_markdown)
    # ------------------------------------------------------------------
    hmd = cv_markdown.structure_to_markdown(headered["structure"])
    assert "## " in hmd, "markdown is missing ## section headings"
    assert "OBSERVED LAYOUT" in hmd, "markdown is missing the OBSERVED LAYOUT subsection"
    assert "(page " in hmd, "markdown layout must carry (page N, y approx Y) markers"
    assert "## EXPERIENCE" in hmd and "## SKILLS" in hmd, hmd

    md_docx = cv_markdown.structure_to_markdown(docx_parsed["structure"])
    assert "## " in md_docx
    assert "## WORK EXPERIENCE" in md_docx and "## SKILLS" in md_docx, md_docx
    assert "OBSERVED LAYOUT" in md_docx
    assert "(paragraph " in md_docx, "docx layout must carry (paragraph N) markers"

    # dispatch: pdf-structure -> layout-aware renderer.
    md_pdf = cv_markdown.document_markdown(headered["cv"], headered["text"], headered["structure"])
    assert md_pdf and "## EXPERIENCE" in md_pdf
    # dispatch: docx-structure -> docx renderer.
    md_dx = cv_markdown.document_markdown(docx_parsed["cv"], docx_parsed["text"], docx_parsed["structure"])
    assert "## WORK EXPERIENCE" in md_dx
    # dispatch: no structure -> cv_to_markdown fallback (## headings still present).
    md_none = cv_markdown.document_markdown(make_cv(), "x", None)
    assert md_none and "## EXPERIENCE" in md_none

    # ------------------------------------------------------------------
    # S7 -- CapturingLLM: analyze JD mode + generic mode + vision blocks
    # ------------------------------------------------------------------
    jd_llm = CapturingLLM(CANNED_ANALYSIS)
    jd_rep = cv_analyzer.analyze(make_cv(), "x", "python k8s", jd_llm, structure=headered["structure"])
    assert isinstance(jd_llm.user_content, str), type(jd_llm.user_content)
    assert "Target Job Description:" in jd_llm.user_content
    assert "python k8s" in jd_llm.user_content
    assert "## " in jd_llm.user_content and "OBSERVED LAYOUT" in jd_llm.user_content
    assert jd_rep.ats_score == 82
    assert jd_rep.missing_keywords[0].keyword == "kubernetes"

    gen_llm = CapturingLLM(CANNED_GENERIC_ANALYSIS)
    gen_rep = cv_analyzer.analyze(make_cv(), "x", "", gen_llm, structure=headered["structure"])
    assert isinstance(gen_llm.user_content, str), type(gen_llm.user_content)
    assert "GENERIC ATS-friendliness" in gen_llm.user_content, "generic prompt marker missing"
    assert "exact shape" in gen_llm.user_content, "JSON schema marker missing from prompt"
    assert "## " in gen_llm.user_content and "OBSERVED LAYOUT" in gen_llm.user_content
    assert gen_rep.gaps, f"generic-mode report must carry gaps, got {gen_rep.gaps}"
    assert gen_rep.matched_keywords == [] and gen_rep.missing_keywords == [], "generic mode must not match keywords"

    # vision path: user content is a canonical block list (text first, images after).
    vision_llm = CapturingLLM(CANNED_ANALYSIS)
    vision_rep = cv_analyzer.analyze(
        scan["cv"], scan["text"], "python aws k8s", vision_llm,
        images=scan["page_images"], structure=scan["structure"],
    )
    assert vision_rep.ats_score == 82
    content = vision_llm.user_content
    assert isinstance(content, list) and content, f"expected block list, got {type(content)}"
    assert content[0]["type"] == "text"
    image_blocks = [b for b in content if b["type"] == "image"]
    assert len(image_blocks) == 2, f"expected 2 image blocks, got {len(image_blocks)}"
    first_img = image_blocks[0]
    assert first_img.get("format") == "png"
    raw = base64.b64decode(first_img["bytes"])
    assert raw[:4] == b"\x89PNG", f"not a PNG, magic={raw[:8]!r}"
    joined = " ".join(b.get("text", "") for b in content).lower()
    for kw in ("transcribe", "layout", "pages"):
        assert kw in joined, f"vision prompt missing keyword {kw!r}"

    # text-PDF suggest path: plain-string markdown carrying ## + OBSERVED LAYOUT.
    tp = parse_file(make_text_pdf(), "resume.pdf")
    sug_llm_text = CapturingLLM(CANNED_SUGGESTIONS)
    tsugg = cv_tailor.generate_suggestions(
        tp["cv"], tp["text"], "python java aws", sug_llm_text, structure=tp["structure"]
    )
    t_content = getattr(sug_llm_text, "user_content")
    assert isinstance(t_content, str), f"expected plain string, got {type(t_content)}"
    assert "## " in t_content and "OBSERVED LAYOUT" in t_content
    assert tsugg and tsugg[0].type == "reword" and tsugg[0].field_path, tsugg

    # generic suggest with an empty JD must fabricate nothing.
    assert cv_tailor._keyword_suggestions(make_cv(), "x", "") == []
    empty_sug_llm = CapturingLLM('{"suggestions": []}')
    generic_suggs = cv_tailor.generate_suggestions(make_cv(), "x", "", empty_sug_llm)
    assert generic_suggs == [], f"expected no suggestions for empty JD, got {generic_suggs}"

    # ------------------------------------------------------------------
    # S8 -- field_path application semantics (no LLM)
    # ------------------------------------------------------------------
    fp_reword = Suggestion(
        id="fp-1", section="experience", field="bullets", type="reword", title="q",
        original="helped build stuff", suggested="led delivery of", reason="r",
        field_path="experience[0].bullets[0]",
    )
    applied = cv_tailor.apply_suggestion(make_cv(), fp_reword)
    assert applied.experience[0].bullets == ["led delivery of"], applied.experience[0].bullets

    fp_add = Suggestion(
        id="fp-2", section="custom_sections", field="bullets", type="add", title="q",
        suggested="Volunteer: coordinated donation drives", reason="r",
        field_path="custom_sections[0].bullets",
    )
    custom_cv = make_cv_with_custom_sections()
    custom_applied = cv_tailor.apply_suggestion(custom_cv, fp_add)
    assert "Volunteer: coordinated donation drives" in custom_applied.custom_sections[0].bullets

    fp_bad = Suggestion(
        id="fp-3", section="experience", field="bullets", type="reword", title="q",
        original="x", suggested="y", reason="r", field_path="experience[x].bullets",
    )
    unchanged = cv_tailor.apply_suggestion(make_cv(), fp_bad)  # malformed path: no crash
    assert unchanged == make_cv(), "malformed field_path must leave the CV unchanged"

    # rename-type suggestion: parses with field_path and applies without crashing.
    ren_llm = CapturingLLM(CANNED_SUGGESTIONS)
    ren_suggs = cv_tailor.generate_suggestions(make_cv(), "x", "python aws", ren_llm)
    ren = [s for s in ren_suggs if s.type == "rename"]
    assert ren and ren[0].field_path, ren_suggs
    ren_applied = cv_tailor.apply_suggestion(make_cv(), ren[0])  # must not crash
    assert ren_applied is not None

    # ------------------------------------------------------------------
    # S9 -- deterministic rendering: HTML -> PDF (fitz.Story) + DOCX (no AI)
    # ------------------------------------------------------------------
    html = render_html(make_cv(), "modern")
    assert "Ada Lovelace" in html
    esc_html = render_html(CVData(personal={"name": "Ada"}, summary="A & B"), "modern")
    assert "&amp;" in esc_html and "A &amp; B" in esc_html, "user content must be HTML-escaped"

    pdf_bytes = html_to_pdf(html)
    assert pdf_bytes.startswith(b"%PDF"), pdf_bytes[:8]
    rp = render_pdf(make_cv(), "modern")
    assert rp.startswith(b"%PDF")
    doc = fitz.open(stream=rp, filetype="pdf")
    assert doc.page_count >= 1, "modern render must paginate to >= 1 page"
    doc.close()

    long_pdf = render_pdf(make_long_cv(), "modern")
    assert long_pdf.startswith(b"%PDF")
    doc = fitz.open(stream=long_pdf, filetype="pdf")
    assert doc.page_count > 1, f"long CV must be multi-page, got {doc.page_count}"
    doc.close()

    dd = export_pdf(make_cv(), "deedy-resume")
    assert dd.startswith(b"%PDF")
    doc = fitz.open(stream=dd, filetype="pdf")
    dd_text = " ".join(page.get_text() for page in doc)
    doc.close()
    assert "Ada Lovelace" in dd_text, dd_text[:200]

    # unknown template id falls back to the builtin renderer (still a valid PDF).
    unknown_pdf = render_pdf(make_cv(), "no-such-template")
    assert unknown_pdf.startswith(b"%PDF")
    doc = fitz.open(stream=unknown_pdf, filetype="pdf")
    assert doc.page_count >= 1
    doc.close()

    # DOCX export must render custom_sections (heading + bullets).
    d = export_docx(make_cv())
    assert isinstance(d, bytes) and len(d) > 500
    d_custom = export_docx(make_cv_with_custom_sections())
    re_doc = docx.Document(io.BytesIO(d_custom))
    d_texts = [p.text for p in re_doc.paragraphs]
    assert "VOLUNTEERING" in d_texts, d_texts
    assert any("Helped at a shelter" in t for t in d_texts), d_texts

    # round-trip: canonical CV -> PDF -> parse -> text keeps name + a bullet.
    rt_pdf = render_pdf(make_cv(), "modern")
    rt = parse_file(rt_pdf, "rt.pdf")
    assert rt["is_image_pdf"] is False
    assert "Ada Lovelace" in rt["text"], rt["text"][:200]
    assert ("helped build" in rt["text"]) or ("stuff" in rt["text"]), rt["text"][:400]

    # ------------------------------------------------------------------
    # S10 -- gap engine (deterministic rules, NO AI)
    # ------------------------------------------------------------------
    sketchy_cv = CVData()
    sketchy_gaps = cv_gaps.deterministic_checks(sketchy_cv, None, "no contact no summary")
    assert sketchy_gaps, "sketchy CV must produce deterministic gaps"
    assert any(g.field_path == "personal.email" for g in sketchy_gaps), sketchy_gaps
    assert any(g.field_path == "summary" for g in sketchy_gaps), sketchy_gaps

    rich_cv = make_cv()
    rich_cv.experience = [
        ExperienceItem(company="Acme", role="Engineer", bullets=["did some stuff", "worked hard"])
    ]
    rich_gaps = cv_gaps.deterministic_checks(rich_cv, None, "rich")
    assert any(
        g.issue == "No quantified outcome" and g.field_path == "experience[0].bullets[0]"
        for g in rich_gaps
    ), rich_gaps

    # raising client: mode generic, no jd_profile, deterministic only.
    result = cv_gaps.gap_analysis(make_cv(), "x", None, "", client=RaisingLLM())
    assert result["mode"] == "generic"
    assert result["jd_profile"] is None
    assert isinstance(result["deterministic"], list) and result["deterministic"]
    assert result["semantic"] == [], result["semantic"]

    # working FakeLLM: jd mode populates jd_profile + semantic gaps.
    result_jd = cv_gaps.gap_analysis(make_cv(), "x", None, "python k8s kubernetes", client=fake)
    assert result_jd["mode"] == "jd"
    assert result_jd["jd_profile"] is not None
    assert result_jd["jd_profile"].role_title == "Senior Python Engineer"
    assert result_jd["semantic"] and result_jd["semantic"][0].kind == "semantic"
    assert result_jd["semantic"][0].field_path == "experience[0].bullets[0]"

    # POST /api/cv/gaps (no LLM configured -> deterministic only, generic mode).
    r = client.post("/api/cv/gaps", json={"cv": CVData().model_dump()})
    assert r.status_code == 200, r.text
    gaps_body = r.json()
    assert gaps_body["mode"] == "generic"
    assert gaps_body["jd_profile"] is None
    assert gaps_body["deterministic"], gaps_body
    assert gaps_body["semantic"] == []

    # ------------------------------------------------------------------
    # S11 -- JD parser (LLM canned + deterministic keyword fallback)
    # ------------------------------------------------------------------
    prof = cv_jd_parser.parse_jd("Senior Python Engineer: 5 years Python, Docker, Kubernetes.", fake)
    assert prof is not None
    assert prof.role_title == "Senior Python Engineer"
    assert "Python" in prof.required_skills and "Kubernetes" in prof.required_skills
    assert "senior" in prof.seniority_signals

    prof2 = cv_jd_parser.parse_jd(
        "Senior Software Engineer - python, k8s, 5+ years experience, cloud.",
        CapturingLLM("this is not json at all"),
    )
    assert prof2 is not None
    assert "python" in prof2.must_have_keywords, prof2.must_have_keywords
    assert "k8s" in prof2.must_have_keywords, prof2.must_have_keywords

    assert cv_jd_parser.parse_jd("", fake) is None

    # ------------------------------------------------------------------
    # S12 -- session store (bounds 25, TTL, round trip, flags/classification)
    # ------------------------------------------------------------------
    assert cv_session.session_stats()["count"] == 0, "store should be empty before S12"

    s_a = cv_session.create_session(
        "pdf", "a.pdf", b"orig-data", "Hello Ada",
        {"kind": "pdf"}, {"personal": {"name": "Ada"}}, 0.5, False, [], 0,
    )
    assert cv_session.get_session(s_a) is not None
    assert cv_session.get_session(s_a).text == "Hello Ada"
    assert cv_session.get_session("does-not-exist-xyz") is None
    assert cv_session.delete_session(s_a) is True
    assert cv_session.get_session(s_a) is None

    # create_session carries confidence_flags + classification through the store.
    s_flags = cv_session.create_session(
        "pdf", "flags.pdf", b"x", "t", {"kind": "pdf"}, {"personal": {}}, 0.5, False, [], 0,
        confidence_flags=[{"field_path": "personal.email", "level": "low", "reason": "nd"}],
        classification="heuristic",
    )
    entry = cv_session.get_session(s_flags)
    assert entry is not None and entry.confidence_flags[0]["field_path"] == "personal.email"
    assert entry.classification == "heuristic"
    assert cv_session.delete_session(s_flags) is True

    # 26 tiny sessions -> bounded to MAX_SESSIONS (25).
    created = []
    for i in range(26):
        created.append(
            cv_session.create_session(
                "pdf", f"tiny{i}.pdf", b"x", f"text {i}",
                {"kind": "pdf"}, {"personal": {}}, 0.5, False, [], 0,
            )
        )
    stats = cv_session.session_stats()
    assert stats["count"] == 25, f"expected bound of 25, got {stats}"

    # TTL expiry: force-expire by backdating `created`.
    s_ttl = cv_session.create_session(
        "pdf", "ttl.pdf", b"y", "z", {"kind": "pdf"}, {"personal": {}}, 0.5, False, [], 0,
    )
    created.append(s_ttl)
    entry = cv_session.get_session(s_ttl)
    assert entry is not None
    entry.created -= 2000  # push past DEFAULT_TTL (1800s)
    assert cv_session.get_session(s_ttl) is None

    for sid in created:
        cv_session.delete_session(sid)
    assert cv_session.session_stats()["count"] == 0

    # ------------------------------------------------------------------
    # S13 -- library (round trip + 404)
    # ------------------------------------------------------------------
    r = client.post("/api/library", json={"name": "Ada CV", "cv": make_cv().model_dump()})
    cid = r.json()["id"]
    assert client.get("/api/library").status_code == 200
    got = client.get(f"/api/library/{cid}").json()
    assert got["cv"]["personal"]["name"] == "Ada Lovelace"
    assert client.delete(f"/api/library/{cid}").status_code == 200
    assert client.get(f"/api/library/{cid}").status_code == 404

    # ------------------------------------------------------------------
    # S14 -- analyze + suggest through the API (FakeLLM override)
    # ------------------------------------------------------------------
    # analyze JD mode.
    r = client.post(
        "/api/cv/analyze",
        json={"cv": make_cv().model_dump(), "text": "x", "job_description": "python k8s"},
    )
    assert r.status_code == 200, r.text
    rep = r.json()
    assert rep["ats_score"] == 82
    assert rep["missing_keywords"][0]["keyword"] == "kubernetes"

    # analyze generic mode (no job_description key) -> gaps still present.
    gen_payload = make_cv().model_dump()
    r = client.post("/api/cv/analyze", json={"cv": gen_payload, "text": "x"})
    assert r.status_code == 200, r.text
    grep_rep = r.json()
    assert grep_rep.get("gaps"), f"expected non-empty gaps, got {grep_rep.get('gaps')}"
    assert "ats_score" in grep_rep
    r = client.post("/api/cv/analyze", json={"cv": gen_payload, "text": "x", "job_description": ""})
    assert r.status_code == 200, r.text
    assert r.json().get("gaps"), r.text

    # suggest JD mode -> reword with a usable field_path.
    r = client.post(
        "/api/cv/tailor/suggest",
        json={"cv": make_cv().model_dump(), "text": "x", "job_description": "jd"},
    )
    assert r.status_code == 200, r.text
    sugg = r.json()["suggestions"][0]
    assert sugg["type"] == "reword"
    assert sugg["field_path"], f"suggestion must carry field_path, got {sugg!r}"

    # apply suggestions through the /api/cv/tailor/apply endpoint.
    r = client.post(
        "/api/cv/tailor/apply",
        json={"cv": make_cv().model_dump(), "suggestions": [sugg]},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert "cv" in body and "applied" in body
    assert body["applied"] == 1
    assert body["cv"]["experience"][0]["bullets"] == ["led delivery of"]

    # ------------------------------------------------------------------
    # S15 -- assist kinds + bogus-kind 400
    # ------------------------------------------------------------------
    assist_body = {"cv": make_cv().model_dump(), "content": "solved hard problems at work"}
    for kind in ("optimize", "optimize_summary", "optimize_bullets", "summary", "bullets"):
        r = client.post("/api/cv/assist", json={**assist_body, "kind": kind})
        assert r.status_code == 200, (kind, r.text)
        assert isinstance(r.json().get("text"), str) and r.json()["text"], (kind, r.json())
    r = client.post("/api/cv/assist", json={**assist_body, "kind": "bogus-kind"})
    assert r.status_code == 400, r.text

    # ------------------------------------------------------------------
    # S16 -- endpoint exports (parse + from-scratch render exports)
    # ------------------------------------------------------------------
    # DOCX parses end-to-end with the canonical response contract.
    r = client.post(
        "/api/cv/parse",
        files={
            "file": (
                "resume.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200, r.text
    dbody = r.json()
    assert isinstance(dbody["session_id"], str) and len(dbody["session_id"]) == 32, dbody["session_id"]
    assert dbody["classification"] == "heuristic", "no LLM in tests => heuristic classification"
    assert dbody["confidence_flags"], "confidence_flags must be non-empty for heuristic parse"
    assert dbody["image_mode"] is False
    created_session_ids.append(dbody["session_id"])

    # PDF parses end-to-end.
    r = client.post("/api/cv/parse", files={"file": ("resume.pdf", make_text_pdf(), "application/pdf")})
    assert r.status_code == 200, r.text
    pbody = r.json()
    sid = pbody["session_id"]
    assert isinstance(sid, str) and len(sid) == 32, sid
    assert pbody["classification"] == "heuristic"
    assert pbody["confidence_flags"] and pbody["confidence_flags"][0]["field_path"]
    assert pbody["image_mode"] is False
    created_session_ids.append(sid)

    # the session store carries the same classification + flags.
    sess_entry = cv_session.get_session(sid)
    assert sess_entry is not None
    assert sess_entry.classification == "heuristic"
    assert sess_entry.confidence_flags, sess_entry

    # scanned PDF -> image_mode True.
    r = client.post("/api/cv/parse", files={"file": ("scan.pdf", make_scanned_pdf(), "application/pdf")})
    assert r.status_code == 200, r.text
    sbody = r.json()
    assert sbody["image_mode"] is True
    assert sbody["classification"] == "heuristic"
    assert sbody["confidence_flags"]
    created_session_ids.append(sbody["session_id"])

    # non-PDF/non-DOCX extension is still rejected at the endpoint.
    r = client.post("/api/cv/parse", files={"file": ("notes.txt", b"whatever", "text/plain")})
    assert r.status_code == 400, r.text
    assert "PDF or DOCX" in r.json().get("detail", ""), r.text

    # analyze bound to a parsed session still works.
    r = client.post(
        "/api/cv/analyze",
        json={"session_id": sid, "cv": make_cv().model_dump(), "text": "x", "job_description": "python k8s"},
    )
    assert r.status_code == 200, r.text
    assert r.json()["ats_score"] == 82

    # /api/export/pdf with an explicit template -> render-mode PDF containing name.
    r = client.post("/api/export/pdf", json={"cv": make_cv().model_dump(), "template": "awesome-cv"})
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith("application/pdf"), r.headers.get("content-type")
    ep_pdf = fitz.open(stream=r.content, filetype="pdf")
    ep_text = " ".join(page.get_text() for page in ep_pdf)
    ep_pdf.close()
    assert "Ada Lovelace" in ep_text, ep_text[:200]

    # /api/export/docx -> docx containing name.
    r = client.post("/api/export/docx", json={"cv": make_cv().model_dump()})
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ), r.headers.get("content-type")
    ep_doc = docx.Document(io.BytesIO(r.content))
    assert any("Ada Lovelace" in p.text for p in ep_doc.paragraphs)

    # preserved export is gone: NOT 200 (404 or 405).
    r = client.post("/api/export/preserved", json={"session_id": "bogus", "suggestions": []})
    assert r.status_code in (404, 405), r.status_code

    # gallery-template exports (render-mode; no fill-warnings header).
    r = client.post("/api/templates/nope/export/pdf", json={"cv": make_cv().model_dump()})
    assert r.status_code == 404, r.text
    r = client.post("/api/templates/awesome-cv/export/pdf", json={"cv": make_cv().model_dump()})
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith("application/pdf"), r.headers.get("content-type")
    assert "X-CVIQ-Fill-Warnings" not in r.headers, "render-mode export must not carry fill warnings"
    tpl_pdf = fitz.open(stream=r.content, filetype="pdf")
    tpl_text = " ".join(page.get_text() for page in tpl_pdf)
    tpl_pdf.close()
    assert "Ada Lovelace" in tpl_text, tpl_text[:200]
    r = client.post("/api/templates/cvresume/export/docx", json={"cv": make_cv().model_dump()})
    assert r.status_code == 200, r.text
    tpl_doc = docx.Document(io.BytesIO(r.content))
    assert any("Ada Lovelace" in p.text for p in tpl_doc.paragraphs)

    # ------------------------------------------------------------------
    # S17 -- unified template catalog (3 builtin + 8 gallery, no slot detection)
    # ------------------------------------------------------------------
    r = client.get("/api/templates")
    assert r.status_code == 200, r.text
    all_templates = r.json()["templates"]
    assert len(all_templates) == 11, f"expected 11 templates (3 builtin + 8 gallery), got {len(all_templates)}"
    builtin = [t for t in all_templates if t.get("source") == "builtin"]
    gallery = [t for t in all_templates if t.get("source") == "gallery"]
    assert [t["id"] for t in builtin] == ["modern", "classic", "minimal"], builtin
    assert all(t.get("converted") is True and t.get("pages") == 0 and t.get("preview_url") == "" for t in builtin)
    assert all(t.get("render_template") == t["id"] for t in builtin), builtin
    assert len(gallery) == 8, f"expected 8 gallery templates, got {len(gallery)}"
    ids = [tpl["id"] for tpl in gallery]
    assert all(tpl.get("render_template") for tpl in gallery), "render_template missing in metadata"
    assert all(tpl.get("render_template") == tpl["id"] for tpl in gallery), "gallery render_template must be identity"
    assert all("converted" in tpl for tpl in gallery), "converted missing in metadata"
    assert all(tpl.get("converted") for tpl in gallery), "all gallery designs must be converted"
    assert len(set(ids)) == len(ids), f"template ids not unique: {ids}"
    bad_url = [i for i in ids if not re.fullmatch(r"[a-z0-9-]+", i)]
    assert not bad_url, f"template ids not URL-safe: {bad_url}"
    dead = [tpl["id"] for tpl in gallery if not (template_service.template_bytes(tpl["id"]) or b"")]
    assert not dead, f"templates with no loadable bytes: {dead}"

    r = client.get("/api/templates/awesome-cv/preview/0")
    assert r.status_code == 200, r.text
    assert r.content[:4] == b"\x89PNG", f"not a PNG: {r.content[:8]!r}"
    for tpl_id in ids:
        r = client.get(f"/api/templates/{tpl_id}/preview/0")
        assert r.status_code == 200, (tpl_id, r.text)
        assert r.content[:4] == b"\x89PNG", (tpl_id, r.content[:8])

    # ------------------------------------------------------------------
    # S19 -- structural suggestions: reorder (item + section) and rename
    # ------------------------------------------------------------------
    two_job = make_cv()
    two_job.experience = [
        ExperienceItem(company="Acme", role="Engineer", bullets=["a"]),
        ExperienceItem(company="Globex", role="Engineer", bullets=["b"]),
    ]
    reorder_item = Suggestion(
        id="re-1", section="experience", field="", type="reorder", title="Move item",
        original="", suggested="", reason="r", move_from="0", move_to="1",
    )
    moved, applied_item = cv_tailor.apply_suggestions(two_job, [reorder_item])
    assert applied_item == 1, f"item reorder must apply, got {applied_item}"
    assert moved.experience[0].company == "Globex", moved.experience
    assert moved.experience[1].company == "Acme", moved.experience

    reorder_sec = Suggestion(
        id="re-2", section="skills", field="", type="reorder", title="Move skills up",
        original="", suggested="", reason="r", target_section="skills", move_to="top",
    )
    sec_moved, applied_sec = cv_tailor.apply_suggestions(make_cv(), [reorder_sec])
    assert applied_sec == 1, f"section reorder must apply, got {applied_sec}"
    assert sec_moved.section_order and sec_moved.section_order[0] == "skills", sec_moved.section_order
    sec_html = render_html(sec_moved, "modern")
    assert sec_html.find("<h2>SKILLS</h2>") < sec_html.find("<h2>EXPERIENCE</h2>"), "Skills must render before Experience"

    rename_std = Suggestion(
        id="rn-1", section="experience", field="", type="rename", title="Rename",
        original="Experience", suggested="Work History", reason="r",
    )
    renamed, applied_rn = cv_tailor.apply_suggestions(make_cv(), [rename_std])
    assert applied_rn == 1, f"rename must apply, got {applied_rn}"
    assert renamed.section_titles.get("experience") == "Work History", renamed.section_titles
    assert "<h2>Work History</h2>" in render_html(renamed, "modern")

    custom_ren = make_cv_with_custom_sections()
    rename_custom = Suggestion(
        id="rn-2", section="custom_sections", field="title", index=0, type="rename",
        title="Rename", original="Volunteering", suggested="Community", reason="r",
    )
    custom_ren2, applied_rc = cv_tailor.apply_suggestions(custom_ren, [rename_custom])
    assert applied_rc == 1, f"custom rename must apply, got {applied_rc}"
    assert custom_ren2.custom_sections[0].title == "Community", custom_ren2.custom_sections

    # ------------------------------------------------------------------
    # S20 -- idempotent apply: same add-bullet set twice -> applied==0, no dup
    # ------------------------------------------------------------------
    add_bullet = Suggestion(
        id="idem-1", section="experience", field="bullets", index=0, type="add",
        title="Add bullet", original="", suggested="Led delivery of widgets", reason="r",
    )
    idem_cv, applied_a = cv_tailor.apply_suggestions(make_cv(), [add_bullet])
    assert applied_a == 1, f"first add must apply, got {applied_a}"
    assert idem_cv.experience[0].bullets.count("Led delivery of widgets") == 1
    idem_cv2, applied_b = cv_tailor.apply_suggestions(idem_cv, [add_bullet])
    assert applied_b == 0, f"re-apply must be idempotent, got {applied_b}"
    assert idem_cv2 == idem_cv, "re-applied CV must be unchanged"
    assert idem_cv2.experience[0].bullets.count("Led delivery of widgets") == 1

    # ------------------------------------------------------------------
    # S21 -- chat 2.0: strict JSON reply + proposed_edits; expired session 200
    # ------------------------------------------------------------------
    chat_llm = CapturingLLM(CANNED_CHAT)
    app.dependency_overrides[cv_routes.get_llm_client] = lambda: chat_llm
    r = client.post(
        "/api/cv/tailor/chat",
        json={
            "cv": make_cv().model_dump(),
            "messages": [{"role": "user", "content": "Make my summary stronger"}],
            "job_description": "python",
        },
    )
    assert r.status_code == 200, r.text
    chat_body = r.json()
    assert chat_body["reply"] == "I tightened your summary to lead with impact.", chat_body
    assert chat_body["proposed_edits"] and chat_body["proposed_edits"][0]["type"] == "rewrite", chat_body
    assert chat_body["proposed_edits"][0]["field_path"] == "summary", chat_body
    assert chat_body["session_warning"] == "", chat_body

    # expired session -> 200 (NOT 404) with a structured session_warning.
    r = client.post(
        "/api/cv/tailor/chat",
        json={
            "cv": make_cv().model_dump(),
            "messages": [{"role": "user", "content": "hi"}],
            "session_id": "does-not-exist-xyz",
        },
    )
    assert r.status_code == 200, r.text
    assert "session" in r.json()["session_warning"].lower(), r.json()
    app.dependency_overrides[cv_routes.get_llm_client] = lambda: fake

    # ------------------------------------------------------------------
    # S22 -- /api/export/preview (server-rendered HTML, no-store)
    # ------------------------------------------------------------------
    r = client.post("/api/export/preview", json={"cv": make_cv().model_dump(), "template": "classic"})
    assert r.status_code == 200, r.text
    assert r.headers.get("cache-control") == "no-store", r.headers
    pv = r.json()["html"]
    assert "Ada Lovelace" in pv
    assert "Georgia" in pv, "classic template must be honored in the preview"
    r = client.post("/api/export/preview", json={"cv": make_cv().model_dump()})
    assert r.status_code == 200, r.text
    assert "Ada Lovelace" in r.json()["html"]

    # ------------------------------------------------------------------
    # S23 -- unified catalog sources + gallery-id export
    # ------------------------------------------------------------------
    r = client.get("/api/templates")
    tpls = r.json()["templates"]
    assert tpls[0]["source"] == "builtin" and tpls[0]["id"] == "modern", tpls[0]
    assert tpls[1]["source"] == "builtin" and tpls[1]["id"] == "classic", tpls[1]
    assert tpls[2]["source"] == "builtin" and tpls[2]["id"] == "minimal", tpls[2]
    gallery_entries = [t for t in tpls if t["source"] == "gallery"]
    assert gallery_entries and all(t["source"] == "gallery" for t in gallery_entries)

    # /api/export/pdf with a gallery template id resolves via render_template_for.
    r = client.post("/api/export/pdf", json={"cv": make_cv().model_dump(), "template": "newfuture-cv"})
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith("application/pdf"), r.headers.get("content-type")

    # ------------------------------------------------------------------
    # S24 -- library PUT overwrite (id preserved, meta round-trips)
    # ------------------------------------------------------------------
    r = client.post(
        "/api/library",
        json={"name": "Ada CV", "cv": make_cv().model_dump(), "meta": {"template": "modern", "ats_score": 82}},
    )
    cid = r.json()["id"]
    r = client.put(
        f"/api/library/{cid}",
        json={"name": "Ada CV v2", "cv": make_cv().model_dump(), "meta": {"template": "classic", "ats_score": 90}},
    )
    assert r.status_code == 200, r.text
    assert r.json()["id"] == cid, "PUT must preserve the entry id"
    assert r.json()["name"] == "Ada CV v2", r.json()
    assert r.json()["meta"]["ats_score"] == 90, r.json()
    got = client.get(f"/api/library/{cid}").json()
    assert got["id"] == cid and got["name"] == "Ada CV v2", got
    assert got["meta"]["template"] == "classic", got
    lst = client.get("/api/library").json()
    entry = next(e for e in lst if e["id"] == cid)
    assert entry["meta"]["ats_score"] == 90, entry
    assert client.put("/api/library/nope", json={"cv": make_cv().model_dump()}).status_code == 404
    assert client.delete(f"/api/library/{cid}").status_code == 200

    # ------------------------------------------------------------------
    # S25 -- upload size guard (>15 MB -> 413)
    # ------------------------------------------------------------------
    big_bytes = b"x" * (15 * 1024 * 1024 + 1)
    r = client.post("/api/cv/parse", files={"file": ("big.pdf", big_bytes, "application/pdf")})
    assert r.status_code == 413, r.text
    assert "15 MB" in r.json().get("detail", ""), r.text

    # ------------------------------------------------------------------
    # S26 -- /api/health
    # ------------------------------------------------------------------
    r = client.get("/api/health")
    assert r.status_code == 200, r.text
    assert r.json() == {"status": "ok"}, r.json()

    # ------------------------------------------------------------------
    # S27 -- section_order honored by >=2 templates + the fallback renderer
    # ------------------------------------------------------------------
    ordered = make_cv()
    ordered.section_order = [
        "skills", "summary", "experience", "education", "projects",
        "certifications", "languages", "custom_sections",
    ]
    classic_html = render_html(ordered, "classic")
    assert classic_html.find("<h2>Skills</h2>") < classic_html.find("<h2>Experience</h2>"), "classic must honor section_order"
    awesome_html = render_html(ordered, "awesome-cv")
    assert awesome_html.find("<h2>SKILLS</h2>") < awesome_html.find("<h2>EXPERIENCE</h2>"), "awesome-cv must honor section_order"
    fallback_html = render_html(ordered, "no-such-template")
    assert fallback_html.find("<h2>Skills</h2>") < fallback_html.find("<h2>Experience</h2>"), "fallback renderer must honor section_order"

    # ------------------------------------------------------------------
    # S18 -- config save + redaction through the API (last: it configures the
    # LLM, so nothing LLM-dependent may run after it).
    # ------------------------------------------------------------------
    r = client.post(
        "/api/config",
        json={"provider": "openrouter", "openrouter_model": "test/model", "openrouter_api_key": "secret123"},
    )
    assert r.status_code == 200, r.text
    r = client.get("/api/config")
    cfg = r.json()
    assert cfg["openrouter_api_key"] == "***", "secrets must be redacted when read back"
    assert cfg["configured"] is True
    # redacted placeholder preserves the saved key.
    r = client.post("/api/config", json={"provider": "openrouter", "openrouter_api_key": "***"})
    assert r.status_code == 200, r.text
    r = client.get("/api/config")
    assert r.json()["openrouter_api_key"] == "***"

    # ------------------------------------------------------------------
    # cleanup
    # ------------------------------------------------------------------
    for sid in created_session_ids:
        cv_session.delete_session(sid)
    app.dependency_overrides.clear()
    print("ALL SMOKE TESTS PASSED")


if __name__ == "__main__":
    test_smoke()