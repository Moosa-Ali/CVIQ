"""PDF/DOCX CV parsing into structured CVData.

Both PDF and DOCX/DOC uploads are accepted (:func:`parse_file` rejects every other
extension, e.g. TXT). The PDF parser extracts text and layout blocks with PyMuPDF,
detects section headings using both string aliases and font/size/position
heuristics, and converts the result into a structured :class:`CVData` object plus a
layout ``structure`` dict used by the AI pipeline (markdown rendering, layout
analysis). The DOCX parser uses python-docx to extract paragraphs (with style/run
metadata) and tables into a ``kind == "docx"`` structure dict.
"""

import io
import re
from collections import Counter

from .models import (
    CertificationItem,
    CVData,
    CustomSection,
    DateRange,
    EducationItem,
    ExperienceItem,
    LanguageItem,
    PersonalInfo,
    ProjectItem,
    SkillGroup,
)

_SECTION_ALIASES = {
    "summary": {"summary", "profile", "professional summary", "objective", "about"},
    "experience": {"experience", "work experience", "employment", "work history", "career history"},
    "education": {"education", "academic background"},
    "skills": {"skills", "technical skills", "core competencies", "competencies", "technologies"},
    "projects": {"projects", "project experience", "select projects"},
    "certifications": {"certifications", "certificates", "licenses", "certification"},
    "languages": {"languages", "language"},
    "personal": {"contact", "personal details", "contact information"},
}

_SECTION_ORDER = ["summary", "experience", "education", "skills", "projects", "certifications", "languages"]

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(?:\+?\d[\d\s().-]{6,}\d)")
_URL_RE = re.compile(r"(?:https?://|www\.)[\w./?#=&%:-]+")
_LINKEDIN_RE = re.compile(r"(?:linkedin\.com|github\.com)/[\w./-]+")
_HEADING_RE = re.compile(r"^[A-Z][A-Z0-9 /&.-]{2,}$")


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    def add(text: str):
        if text and text.strip():
            parts.append(text.strip())

    for paragraph in doc.paragraphs:
        add(paragraph.text)

    # Tables: emit each row's cells as a pipe-joined line (keeps "Role | Company" style).
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # Headers/footers (simple): capture any text they contain.
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            add(paragraph.text)
        for paragraph in section.footer.paragraphs:
            add(paragraph.text)

    return "\n".join(parts)


def _extract_pdf_text_blocks(content: bytes) -> tuple[str, list[dict]]:
    """Extract text from a PDF using PyMuPDF.

    Returns ``(plain_text, blocks)`` where ``blocks`` is a flat list of line dicts:
    ``{"page": n, "text": "<line>", "x0":..., "y0":..., "size":..., "font":...,
    "bold": bool}``. Identical repeated spans are deduplicated.
    """
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    blocks: list[dict] = []
    seen: set[tuple] = set()
    lines_text: list[str] = []
    try:
        for page_num in range(doc.page_count):
            page = doc[page_num]
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                if block.get("type") != 0:  # skip image blocks
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    if not text:
                        continue
                    first = spans[0]
                    size = first.get("size", 0.0)
                    font = first.get("font", "")
                    flags = first.get("flags", 0)
                    bold = bool(flags & 16)  # PyMuPDF flags bit 4 == bold
                    bbox = line.get("bbox", [0.0, 0.0, 0.0, 0.0])
                    key = (text, round(bbox[0], 2), round(bbox[1], 2), round(size, 2), font, bold)
                    if key in seen:
                        continue
                    seen.add(key)
                    blocks.append(
                        {
                            "page": page_num,
                            "text": text,
                            "x0": bbox[0],
                            "y0": bbox[1],
                            "size": size,
                            "font": font,
                            "bold": bold,
                        }
                    )
                    lines_text.append(text)
    finally:
        doc.close()
    return "\n".join(lines_text), blocks


def _extract_docx_structure(content: bytes) -> tuple[str, dict]:
    """Extract structured content from a DOCX using python-docx.

    Returns ``(plain_text, structure)`` where ``structure`` is:
    ``{"kind": "docx", "paragraphs": [...], "tables": [...], "headers_text": [...],
    "footers_text": [...], "sections": [...]}``. The plain text uses the same join
    logic as ``_extract_docx``.
    """
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    def add(text: str):
        if text and text.strip():
            parts.append(text.strip())

    paragraphs: list[dict] = []
    for paragraph in doc.paragraphs:
        runs = []
        for run in paragraph.runs:
            runs.append(
                {
                    "text": run.text,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "underline": bool(run.underline),
                    "size_pt": run.font.size.pt if run.font.size else None,
                }
            )
        alignment = None
        try:
            alignment = paragraph.alignment.name if paragraph.alignment is not None else None
        except Exception:
            alignment = None
        paragraphs.append(
            {
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else None,
                "alignment": alignment,
                "runs": runs,
            }
        )
        add(paragraph.text)

    tables: list[dict] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(cells)
            non_empty = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if non_empty:
                parts.append(" | ".join(non_empty))
        tables.append({"rows": rows})

    headers_text: list[str] = []
    footers_text: list[str] = []
    sections: list[dict] = []
    for section in doc.sections:
        h = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
        f = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
        headers_text.extend(h)
        footers_text.extend(f)
        margins = {}
        try:
            margins = {
                "top": section.top_margin.inches if section.top_margin else None,
                "bottom": section.bottom_margin.inches if section.bottom_margin else None,
                "left": section.left_margin.inches if section.left_margin else None,
                "right": section.right_margin.inches if section.right_margin else None,
            }
        except Exception:
            margins = {}
        sections.append({"margins": margins})
        for p in section.header.paragraphs:
            add(p.text)
        for p in section.footer.paragraphs:
            add(p.text)

    structure = {
        "kind": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "headers_text": headers_text,
        "footers_text": footers_text,
        "sections": sections,
    }
    return "\n".join(parts), structure


def _classify_section(line: str):
    key = line.strip().lower().rstrip(":").rstrip("\\").rstrip(".")
    if len(key) > 40:
        return None
    for canonical, aliases in _SECTION_ALIASES.items():
        if key in aliases:
            return canonical
    return None


def _is_heading(line: str) -> bool:
    """Robust section-heading detection.

    Accepts a line as a heading if its normalized text is a known section alias, OR if
    it is a short capitalized line (<= 30 chars, no sentence-ending period, all words
    capitalized) that matches an alias. Does not require the strict title-case regex.
    """
    stripped = line.strip()
    if not stripped:
        return False
    normalized = stripped.lower().rstrip(":").rstrip("\\").strip()
    if _classify_section(normalized):
        return True
    if len(stripped) <= 30 and not stripped.endswith("."):
        words = stripped.split()
        if words and all(w[:1].isupper() for w in words if w):
            return _classify_section(normalized) is not None
    return False


def _looks_like_pdf_heading(line: str, block: dict, body_size: float) -> bool:
    """Font/size-aware heading heuristic for PDF layout blocks.

    A line is a candidate section heading when it is short (<= 40 chars), does not
    end in a period, is bold OR at least 1.15x the body font size OR all-caps, AND
    its words are mostly capitalized. Used by ``_parse_pdf_text`` and the markdown
    renderer to detect headings that plain-text heuristics would miss.
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 40:
        return False
    if stripped.endswith("."):
        return False
    size = block.get("size", 0.0) or 0.0
    bold = bool(block.get("bold", False))
    all_caps = stripped.isupper() and len(stripped) > 2
    if not (bold or (body_size and size >= body_size * 1.15) or all_caps):
        return False
    words = stripped.split()
    if not words:
        return False
    capitalized = sum(1 for w in words if w and w[0].isupper())
    return capitalized >= len(words) * 0.6


def _parse_pdf_text(text: str, blocks: list[dict]) -> str:
    """Rebuild PDF text with canonical section-heading markers inserted.

    Uses the layout blocks (fonts, sizes, positions) to detect section headings more
    reliably than string heuristics alone. Blocks are sorted in reading order
    (``page``, then ``y0``); the body font size is the most common block size. Each
    detected heading line is replaced by its canonical marker (e.g. ``EXPERIENCE``
    for "Work Experience") so the existing ``_parse_*`` section parsers split
    sections correctly. Falls back to the raw text when no blocks are available.
    """
    if not blocks:
        return text
    ordered = sorted(blocks, key=lambda b: (b.get("page", 0), b.get("y0", 0.0)))
    sizes = [b.get("size", 0.0) for b in ordered if b.get("size", 0.0) > 0]
    body_size = Counter(sizes).most_common(1)[0][0] if sizes else 0.0
    lines: list[str] = []
    for block in ordered:
        line = (block.get("text", "") or "").strip()
        if not line:
            continue
        canonical = _classify_section(line)
        if canonical is None and _looks_like_pdf_heading(line, block, body_size):
            canonical = _classify_section(line)
        lines.append(canonical.upper() if canonical else line)
    return "\n".join(lines)


def _split_sections(lines: list[str]):
    sections: dict[str, list[str]] = {}
    current = None
    for line in lines:
        if _is_heading(line):
            current = _classify_section(line)
            sections.setdefault(current, [])
        elif current:
            sections[current].append(line)
    return sections


def _parse_personal(lines: list[str], text: str) -> PersonalInfo:
    personal = PersonalInfo()
    if lines:
        personal.name = lines[0]
    emails = _EMAIL_RE.findall(text)
    if emails:
        personal.email = emails[0]
    phones = _PHONE_RE.findall(text)
    if phones:
        personal.phone = phones[0]
    urls = _URL_RE.findall(text)
    for url in urls:
        if "linkedin" in url:
            personal.linkedin = url
        elif "github" in url:
            personal.github = url
        elif not personal.location:
            personal.location = url
    return personal


def _fill_experience_header(item: ExperienceItem, header: str):
    """Populate role/company/location from a header line.

    Handles "Role | Company | Location" and "Company — Role" formats.
    """
    header = header.strip()
    if not header:
        return
    if "|" in header:
        parts = [p.strip() for p in header.split("|") if p.strip()]
        if parts:
            item.role = parts[0]
        if len(parts) > 1:
            item.company = parts[1]
        if len(parts) > 2:
            item.location = parts[2]
        return
    dash = re.split(r"\s*[—–]\s*", header)
    if len(dash) == 2 and dash[0].strip() and dash[1].strip():
        item.company = dash[0].strip()
        item.role = dash[1].strip()
        return
    # Single value: prefer role if it reads like a job title, else company.
    if _looks_like_role(header):
        item.role = header
    else:
        item.company = header


def _looks_like_role(text: str) -> bool:
    role_words = {"engineer", "developer", "manager", "analyst", "designer", "lead", "architect", "scientist", "consultant", "specialist", "director", "officer", "intern", "associate", "coordinator"}
    words = text.lower().split()
    return any(w in role_words for w in words)


def _looks_like_header(line: str) -> bool:
    """Heuristic: is this line a role/company header rather than a bullet?"""
    if line.startswith(("-", "•", "*")):
        return False
    if len(line) > 90:
        return False
    if line.endswith((".", "!", "?")):
        return False
    if "|" in line or re.search(r"[—–]", line):
        return True
    words = line.split()
    if len(words) <= 8 and all(w[:1].isupper() for w in words if w):
        return True
    return False


def _parse_experience(raw: list[str]) -> list[ExperienceItem]:
    items: list[ExperienceItem] = []
    item: ExperienceItem | None = None
    for line in raw:
        if not line:
            continue
        if _EMAIL_RE.search(line) or _PHONE_RE.search(line):
            continue
        line = line.strip()

        date_match = _DATE_RE.search(line)
        if date_match:
            header = _DATE_RE.sub("", line).strip(" -|")
            if item is not None and not item.dates.start:
                # Date line belongs to the current item (header-then-date ordering).
                item.dates = DateRange(start=date_match.group(1), end=date_match.group(2))
                if header:
                    _fill_experience_header(item, header)
                continue
            if item is not None:
                items.append(item)
            item = ExperienceItem()
            item.dates = DateRange(start=date_match.group(1), end=date_match.group(2))
            if header:
                _fill_experience_header(item, header)
            continue

        if _looks_like_header(line):
            if item is not None:
                items.append(item)
            item = ExperienceItem()
            _fill_experience_header(item, line)
            continue

        # Bullet line.
        if item is None:
            item = ExperienceItem()
        item.bullets.append(line.lstrip("-•* ").strip())

    if item is not None:
        items.append(item)
    return [i for i in items if i.role or i.company or i.bullets]


_DATE_RE = re.compile(r"([A-Za-z]{3,9}\s?\d{4}|\d{4})[\s-]+([A-Za-z]{3,9}\s?\d{4}|\d{4}|Present|present|Current|current)")


def _parse_education(raw: list[str]) -> list[EducationItem]:
    items: list[EducationItem] = []
    item: EducationItem | None = None
    pending: list[str] = []
    for line in raw:
        if not line:
            continue
        if match := _DATE_RE.search(line):
            if item or pending:
                # NOTE: the previous ``current = ...`` walrus and ``fields`` /
                # ``location_part`` locals were dead code (never read) — removed;
                # behavior is unchanged. A date line finalizes the pending item.
                remaining = _DATE_RE.sub("", line).strip(" -|")
                if item is None:
                    item = EducationItem(institution=pending[0] if pending else "", dates=DateRange(start=match.group(1), end=match.group(2)))
                else:
                    item.dates = DateRange(start=match.group(1), end=match.group(2))
                if remaining and not item.field:
                    item.field = remaining
                items.append(item)
                item = None
                pending = []
            continue
        if item is None:
            pending.append(line.strip())
            if len(pending) == 1:
                item = EducationItem(institution=line.strip())
                pending = []
    if item is not None:
        items.append(item)
    return [i for i in items if i.institution or i.degree or i.field]


def _parse_skills(raw: list[str]) -> list[SkillGroup]:
    groups: list[SkillGroup] = []
    for line in raw:
        line = line.strip()
        if not line:
            continue
        if ":" in line:
            category, rest = line.split(":", 1)
            skills = [s.strip() for s in re.split(r"[,;|]", rest) if s.strip()]
            groups.append(SkillGroup(category=category.strip(), skills=skills))
        else:
            skills = [s.strip() for s in re.split(r"[,;|]", line) if s.strip()]
            groups.append(SkillGroup(skills=skills))
    return groups


def _parse_projects(raw: list[str]) -> list[ProjectItem]:
    projects: list[ProjectItem] = []
    for line in raw:
        if not line.strip():
            continue
        if _HEADING_RE.match(line.strip()) or "|" in line:
            link = ""
            name = line.strip()
            url = _URL_RE.search(line)
            if url:
                link = url.group(0)
                name = name.replace(link, "").strip(" -|()")
            projects.append(ProjectItem(name=name, link=link))
        elif projects:
            projects[-1].bullets.append(line.lstrip("- ").strip())
    return projects


def _parse_certifications(raw: list[str]) -> list[CertificationItem]:
    certs: list[CertificationItem] = []
    for line in raw:
        line = line.strip()
        if not line:
            continue
        certs.append(CertificationItem(name=line))
    return certs


def _parse_languages(raw: list[str]) -> list[LanguageItem]:
    languages: list[LanguageItem] = []
    for line in raw:
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in re.split(r"[,(]", line) if p.strip()]
        if len(parts) > 1:
            languages.append(LanguageItem(name=parts[0], level=parts[1].strip("()")))
        else:
            languages.append(LanguageItem(name=parts[0]))
    return languages


def _parse_text(text: str) -> CVData:
    raw_lines = [line.rstrip() for line in text.splitlines()]
    non_empty = [line for line in raw_lines if line.strip()]
    clean = [line.strip() for line in raw_lines if line.strip()]
    cv = CVData()
    cv.personal = _parse_personal(non_empty, text)
    sections = _split_sections(clean)
    for section in _SECTION_ORDER:
        raw = sections.get(section, [])
        if section == "summary":
            cv.summary = " ".join(raw).strip()
        elif section == "experience":
            cv.experience = _parse_experience(raw)
        elif section == "education":
            cv.education = _parse_education(raw)
        elif section == "skills":
            cv.skills = _parse_skills(raw)
        elif section == "projects":
            cv.projects = _parse_projects(raw)
        elif section == "certifications":
            cv.certifications = _parse_certifications(raw)
        elif section == "languages":
            cv.languages = _parse_languages(raw)
    # Conservative custom-section recovery: non-standard headings that survived
    # the standard heuristic parse are preserved verbatim (title + bullets).
    cv.custom_sections = _parse_custom_sections(clean)
    return cv


def _is_all_caps_heading(line: str) -> bool:
    """A short ALL-CAPS line (``_HEADING_RE`` style, <= 40 chars, no period).

    Used as the "looks like another heading" check for custom headings and as the
    break condition when collecting custom-section bullets — title-cased bullet
    lines deliberately do NOT count so content like "Mentored interns" stays a
    bullet.
    """
    stripped = line.strip()
    return bool(
        stripped
        and len(stripped) <= 40
        and not stripped.endswith(".")
        and _HEADING_RE.match(stripped)
    )


def _is_custom_heading(clean: list[str], i: int) -> bool:
    """Conservative custom-heading test for ``clean[i]``.

    A line is a custom heading when it is NOT a known section (``_classify_section
    is None``) AND it reads like a heading (an ALL-CAPS short line via
    ``_HEADING_RE``, or a title-cased short line whose words are capitalized)
    AND the following non-empty line is not itself another heading. Nothing is
    inferred about content — a heading only becomes a section when at least one
    bullet follows (see :func:`_parse_custom_sections`).
    """
    line = clean[i].strip()
    if not line:
        return False
    if _classify_section(line) is not None:
        return False  # standard heading — handled by the standard loop
    if len(line) > 40 or line.endswith("."):
        return False
    words = line.split()
    if not words:
        return False
    all_caps = bool(_HEADING_RE.match(line))
    title_case = len(words) >= 2 and all(w[:1].isupper() for w in words if w)
    if not (all_caps or title_case):
        return False
    # Conservative guard: the document's FIRST line is almost always the candidate's
    # name (e.g. "Ada Lovelace"), which is a title-cased short line. Only skip it
    # when it is NOT an ALL-CAPS heading, so a leading all-caps custom heading
    # (e.g. "VOLUNTEERING") is still detected.
    if i == 0 and not all_caps:
        return False
    # The following non-empty line must not itself be another heading.
    for nxt in clean[i + 1 :]:
        if not nxt.strip():
            continue
        return not (_is_heading(nxt) or _is_all_caps_heading(nxt))
    return True


def _parse_custom_sections(clean: list[str]) -> list[CustomSection]:
    """Recover NON-standard section headings into ``cv.custom_sections``.

    Scans the cleaned document for custom headings (see :func:`_is_custom_heading`).
    Once one opens, every subsequent line becomes a bullet until the next heading
    (known or custom) or the end of the document. A heading that is followed by no
    bullets is treated as a stray line and ignored — it is never dropped into a
    standard section, so the heuristic cannot corrupt the canonical sections.
    """
    sections: list[CustomSection] = []
    i = 0
    n = len(clean)
    while i < n:
        if not _is_custom_heading(clean, i):
            i += 1
            continue
        bullets: list[str] = []
        j = i + 1
        while j < n and not _is_heading(clean[j]) and not _is_all_caps_heading(clean[j]):
            bullets.append(clean[j])
            j += 1
        if bullets:
            sections.append(CustomSection(title=clean[i], bullets=bullets))
        i = j
    return sections


def _confidence(cv: CVData, text: str) -> float:
    if not text:
        return 0.0
    score = 0.5
    if cv.personal.name:
        score += 0.15
    found = 0
    sections = [bool(cv.experience), bool(cv.education), bool(cv.skills), bool(cv.summary)]
    found = sum(sections)
    score += 0.1 * found
    return max(0.0, min(1.0, score))


def _parse_pdf(content: bytes, filename: str) -> dict:
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    page_count = doc.page_count
    try:
        text, blocks = _extract_pdf_text_blocks(content)
        # Image-PDF detection: little/no extractable text overall, OR any page that
        # yields no text at all (a scanned page).
        is_image_pdf = len(text.strip()) < 40 or (
            page_count > 0
            and any(not (doc[i].get_text() or "").strip() for i in range(page_count))
        )
        page_images: list[bytes] = []
        if is_image_pdf:
            for i in range(page_count):
                pix = doc[i].get_pixmap(dpi=150, colorspace=fitz.csRGB)
                page_images.append(pix.tobytes("png"))
    finally:
        doc.close()

    cv = _parse_text(_parse_pdf_text(text, blocks))
    confidence = 0.0 if is_image_pdf else _confidence(cv, text)
    structure = {"kind": "pdf", "blocks": blocks, "page_count": page_count}
    return {
        "text": text,
        "cv": cv,
        "confidence": confidence,
        "structure": structure,
        "is_image_pdf": is_image_pdf,
        "page_count": page_count,
        "page_images": page_images,
    }


def _parse_docx(content: bytes, filename: str) -> dict:
    text, structure = _extract_docx_structure(content)
    has_paragraphs = any(p.get("text", "").strip() for p in structure["paragraphs"])
    has_tables = any(
        any(cell.strip() for row in t["rows"] for cell in row) for t in structure["tables"]
    )
    if not text.strip() and not has_paragraphs and not has_tables:
        raise ValueError(
            "This DOCX contains no extractable text and is not supported for vision analysis. "
            "Please upload the CV as a text-based DOCX or as a PDF."
        )
    cv = _parse_text(text)
    confidence = _confidence(cv, text)
    return {
        "text": text,
        "cv": cv,
        "confidence": confidence,
        "structure": structure,
        "is_image_pdf": False,
        "page_count": 0,
        "page_images": [],
    }


def parse_file(content: bytes, filename: str) -> dict:
    """Parse an uploaded CV file into a normalized result dict.

    Both PDF and DOCX/DOC uploads are accepted. PDFs are parsed via the
    layout-aware PyMuPDF pipeline (``_parse_pdf``, ``kind == "pdf"``); DOCX/DOC
    files are parsed via the structured python-docx pipeline (``_parse_docx``,
    ``kind == "docx"``). Any other file type (e.g. TXT) is rejected with a
    :class:`ValueError` so the app never ingests an unsupported CV.

    Returns:
        {
            "text": str, "cv": CVData, "confidence": float,
            "structure": dict, "is_image_pdf": bool, "page_count": int,
            "page_images": list[bytes],
        }
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _parse_pdf(content, filename)
    if name.endswith((".docx", ".doc")):
        return _parse_docx(content, filename)
    raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")
