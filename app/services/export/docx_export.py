import io

from ..cv.models import CVData


def _hex_to_rgb(hex_str: str):
    hex_str = (hex_str or "#2563eb").lstrip("#")
    if len(hex_str) != 6:
        hex_str = "2563eb"
    try:
        return tuple(int(hex_str[i : i + 2], 16) / 255.0 for i in (0, 2, 4))
    except ValueError:
        return (0x25 / 255.0, 0x63 / 255.0, 0xEB / 255.0)


def export_docx(cv: CVData, template: str = "", accent: str = "") -> bytes:
    """Canonical CV -> DOCX bytes (deterministic, no AI).

    ``template`` is accepted for API parity with the unified catalog but does
    not change the DOCX layout (the DOCX renderer is a single generic layout).
    ``accent`` overrides the CV's accent color for the name + section headings;
    it defaults to ``cv.accent``.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    r, g, b = _hex_to_rgb(accent or cv.accent)
    accent = RGBColor(int(r * 255), int(g * 255), int(b * 255))

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    personal = cv.personal
    if personal.name:
        name = doc.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name.add_run(personal.name)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = accent

    if personal.title:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(personal.title)
        run.font.size = Pt(13)

    contact = " | ".join(
        part
        for part in [personal.email, personal.phone, personal.location, personal.website, personal.linkedin, personal.github]
        if part
    )
    if contact:
        contact_par = doc.add_paragraph()
        contact_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_par.add_run(contact)

    def add_heading(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.color.rgb = accent
        return p

    if cv.summary:
        add_heading("Summary")
        doc.add_paragraph(cv.summary)

    if cv.experience:
        add_heading("Experience")
        for item in cv.experience:
            header = (item.role + " - " + item.company) if (item.role and item.company) else (item.role or item.company)
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            meta = " | ".join(
                part
                for part in [item.location, f"{item.dates.start} - {item.dates.end}" if (item.dates.start or item.dates.end) else ""]
                if part
            )
            if meta:
                doc.add_paragraph(meta)
            for bullet in item.bullets:
                doc.add_paragraph("\u2022 " + bullet)

    if cv.education:
        add_heading("Education")
        for item in cv.education:
            header = " ".join(x for x in [item.degree, item.field] if x) or item.institution
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            meta = " | ".join(
                part
                for part in [
                    item.institution if header != item.institution else "",
                    item.gpa,
                    f"{item.dates.start} - {item.dates.end}" if (item.dates.start or item.dates.end) else "",
                ]
                if part
            )
            if meta:
                doc.add_paragraph(meta)

    if cv.skills:
        add_heading("Skills")
        for group in cv.skills:
            label = f"{group.category}: " if group.category else ""
            doc.add_paragraph(label + ", ".join(group.skills))

    if cv.projects:
        add_heading("Projects")
        for project in cv.projects:
            header = project.name + (f" ({project.link})" if project.link else "")
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            if project.description:
                doc.add_paragraph(project.description)
            for bullet in project.bullets:
                doc.add_paragraph("\u2022 " + bullet)

    if cv.certifications:
        add_heading("Certifications")
        for cert in cv.certifications:
            suffix = " | ".join(x for x in [cert.issuer, cert.year] if x)
            doc.add_paragraph(cert.name + (f" ({suffix})" if suffix else ""))

    if cv.languages:
        add_heading("Languages")
        for lang in cv.languages:
            doc.add_paragraph(lang.name + (f" ({lang.level})" if lang.level else ""))

    for section in cv.custom_sections:
        if section.title:
            add_heading(section.title)
        for bullet in section.bullets:
            doc.add_paragraph("\u2022 " + bullet)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
